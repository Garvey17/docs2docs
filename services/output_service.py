from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt

from writer.schema import DocSection


OUTPUT_DIR = Path("output")


SECTION_TITLES = {
    "installation": "Installation",
    "getting_started": "Getting Started",
    "features": "Features",
}


def _set_default_font(document: Document) -> None:
    styles = document.styles

    normal_style = styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(12)


def _add_code_paragraph(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()

    run = paragraph.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(10)

    # Light grey background
    paragraph._element.get_or_add_pPr().append(
        parse_xml(
            r'<w:shd {} w:fill="EDEDED"/>'.format(
                nsdecls("w")
            )
        )
    )


def _render_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()

    in_code_block = False
    code_lines: list[str] = []

    for line in lines:
        stripped = line.strip()

        # Code fences
        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False

                code_text = "\n".join(code_lines)
                _add_code_paragraph(document, code_text)

            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Skip empty lines
        if not stripped:
            continue

        # Heading 3
        if stripped.startswith("### "):
            document.add_heading(
                stripped[4:].strip(),
                level=3,
            )
            continue

        # Heading 2
        if stripped.startswith("## "):
            document.add_heading(
                stripped[3:].strip(),
                level=2,
            )
            continue

        # Bullet lists
        if stripped.startswith("- ") or stripped.startswith("* "):
            document.add_paragraph(
                stripped[2:].strip(),
                style="List Bullet",
            )
            continue

        # Numbered lists
        if re.match(r"^\d+\.\s+", stripped):
            text = re.sub(
                r"^\d+\.\s+",
                "",
                stripped,
            )

            document.add_paragraph(
                text,
                style="List Number",
            )
            continue

        # Normal paragraph
        document.add_paragraph(stripped)


def render(
    sections: list[DocSection],
    package_name: str,
) -> Path:
    OUTPUT_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    document = Document()

    _set_default_font(document)

    # Title Page
    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    title_run = title.add_run(package_name)
    title_run.bold = True
    title_run.font.size = Pt(24)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle.add_run(
        f"Documentation Guide — Generated "
        f"{datetime.now():%Y-%m-%d}"
    )

    document.add_page_break()

    # Sections
    for section in sections:
        section_title = SECTION_TITLES.get(
            section.section,
            section.section.replace("_", " ").title(),
        )

        document.add_heading(
            section_title,
            level=1,
        )

        _render_markdown(
            document,
            section.content,
        )

    output_path = (
        OUTPUT_DIR
        / f"{package_name}_guide.docx"
    )

    document.save(output_path)

    return output_path